import os
import re
import json
import time
import uuid
from io import BytesIO
from datetime import datetime, timedelta, timezone
from typing import List, Dict, Optional, Tuple, Any

import pandas as pd
import streamlit as st
import google.generativeai as genai
from supabase import create_client
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont


# =========================================================
# Page Config
# =========================================================
st.set_page_config(
    page_title="Lobster AI",
    page_icon="🦞",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    #MainMenu, footer, header { visibility: hidden; }

    html, body, [class*="css"] {
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
        scroll-behavior: smooth;
    }

    section[data-testid="stSidebar"] {
        display: none !important;
    }

    .block-container {
        padding-top: 0.16rem;
        padding-bottom: 0.20rem;
        padding-left: 0.56rem;
        padding-right: 0.56rem;
        max-width: 1140px;
    }

    .lobster-header { margin-bottom: 0.16rem; }

    .lobster-brand {
        font-size: 1.02rem;
        font-weight: 800;
        line-height: 1.05;
        letter-spacing: -0.02em;
        margin: 0;
        color: #111827;
    }

    .lobster-subtitle {
        font-size: 0.70rem;
        color: #6b7280;
        line-height: 1.15;
        margin-top: 0.02rem;
    }

    .lobster-card {
        border: 1px solid #ececec;
        border-radius: 16px;
        background: #fcfcfd;
        padding: 0.52rem;
        margin-bottom: 0.28rem;
    }

    .lobster-soft-card {
        border: 1px solid #efefef;
        border-radius: 14px;
        background: #ffffff;
        padding: 0.56rem;
        margin-bottom: 0.22rem;
    }

    .lobster-section-title {
        font-size: 0.90rem;
        font-weight: 700;
        margin-bottom: 0.34rem;
        color: #111827;
    }

    .lobster-mini-title {
        font-size: 0.82rem;
        font-weight: 700;
        color: #111827;
        margin-bottom: 0.14rem;
    }

    .lobster-muted {
        font-size: 0.74rem;
        color: #6b7280;
        line-height: 1.35;
    }

    .lobster-note {
        font-size: 0.75rem;
        color: #6b7280;
        margin-top: -0.02rem;
        margin-bottom: 0.12rem;
    }

    .lobster-empty {
        text-align: center;
        padding: 1rem 0.4rem 0.85rem 0.4rem;
        color: #6b7280;
    }

    .lobster-empty h2 {
        margin: 0 0 0.2rem 0;
        font-size: 1rem;
        color: #111827;
    }

    .lobster-pill {
        display: inline-block;
        font-size: 0.68rem;
        padding: 3px 8px;
        border-radius: 999px;
        background: #f3f4f6;
        color: #374151;
        margin-right: 6px;
        margin-bottom: 6px;
    }

    .lobster-task-col {
        border: 1px solid #ececec;
        border-radius: 16px;
        background: #fafafa;
        padding: 0.5rem;
        min-height: 180px;
    }

    .lobster-task-count {
        font-size: 0.72rem;
        color: #6b7280;
        margin-bottom: 0.4rem;
    }

    .lobster-linked-box {
        border: 1px dashed #d8dbe2;
        border-radius: 12px;
        padding: 0.45rem;
        background: #fbfbfc;
        margin-top: 0.35rem;
    }

    .stButton > button, .stDownloadButton > button {
        border-radius: 12px !important;
        min-height: 1.9rem !important;
        padding: 0.22rem 0.66rem !important;
        font-size: 0.84rem !important;
        box-shadow: none !important;
    }

    .stTextInput input,
    .stTextArea textarea,
    .stSelectbox div[data-baseweb="select"] > div,
    .stFileUploader div[data-baseweb="file-uploader"] {
        border-radius: 12px !important;
    }

    div[data-testid="stFileUploader"] section {
        padding: 0.28rem !important;
        border-radius: 12px !important;
    }

    div[data-testid="stExpander"] details {
        border-radius: 12px !important;
        border: 1px solid #ececec !important;
        background: white !important;
    }

    div[data-testid="stExpander"] summary {
        padding-top: 0.35rem !important;
        padding-bottom: 0.35rem !important;
        font-size: 0.88rem !important;
    }

    div[data-testid="stChatMessage"] {
        border-radius: 14px;
        padding-top: 0.02rem !important;
        padding-bottom: 0.02rem !important;
        margin-bottom: 0.08rem !important;
    }

    div[data-testid="stChatMessageContent"] p {
        font-size: 0.97rem;
        line-height: 1.42;
    }

    .element-container { margin-bottom: 0.12rem !important; }

    @media (max-width: 768px) {
        .block-container {
            padding-top: 0.12rem;
            padding-bottom: 0.15rem;
            padding-left: 0.54rem;
            padding-right: 0.54rem;
            max-width: 100%;
        }

        .lobster-brand { font-size: 0.99rem; }
        .lobster-subtitle { font-size: 0.68rem; }

        .lobster-card {
            padding: 0.46rem;
            border-radius: 15px;
        }

        .stButton > button, .stDownloadButton > button {
            min-height: 1.85rem !important;
            font-size: 0.83rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# =========================================================
# ENV
# =========================================================
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")

raw_api_keys = os.getenv("GOOGLE_API_KEYS", "").strip()
single_api_key = os.getenv("GOOGLE_API_KEY", "").strip()

API_KEY_LIST = []
if raw_api_keys:
    API_KEY_LIST = [k.strip() for k in raw_api_keys.split(",") if k.strip()]
elif single_api_key:
    API_KEY_LIST = [single_api_key]

missing = []
if not SUPABASE_URL:
    missing.append("SUPABASE_URL")
if not SUPABASE_ANON_KEY:
    missing.append("SUPABASE_ANON_KEY")
if not API_KEY_LIST:
    missing.append("GOOGLE_API_KEYS or GOOGLE_API_KEY")

if missing:
    st.error("Missing env vars: " + ", ".join(missing))
    st.stop()

supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)


# =========================================================
# Constants
# =========================================================
TASK_TEMPLATES = {
    "Chat": {
        "一般問答": "直接回答使用者問題，先講結論，再補充必要細節。",
        "腦力激盪": "提供多個方向、想法與切入角度，幫助使用者發散思考。",
        "行動建議": "聚焦在下一步該怎麼做，給出可執行建議。",
        "改寫優化": "將使用者內容改寫得更清楚、更有條理、更專業。",
    },
    "Summary": {
        "重點摘要": "先一句總結，再列出 3 到 5 個重點。",
        "條列整理": "把內容整理成條列式重點，方便快速閱讀。",
        "高階主管版摘要": "使用高層決策者視角，強調結論與影響。",
        "可執行建議版": "除了摘要，也補上後續可執行建議。",
    },
    "File Analyze": {
        "結構分析": "說明檔案結構、欄位、內容概覽。",
        "異常檢查": "找出異常、缺值、風險、資料品質問題。",
        "趨勢重點": "找出關鍵趨勢、重要變化與值得注意的地方。",
        "決策建議": "基於檔案內容，整理重點並提出決策建議。",
    },
    "Report": {
        "正式報告": "整理成正式報告格式，適合交付。",
        "會議紀要": "整理成會議紀要，包含背景、重點、待辦。",
        "提案草稿": "整理成提案草稿，強調背景、方案、效益。",
        "進度更新": "整理成進度更新格式，適合對主管或團隊報告。",
    },
    "Memory Assist": {
        "提煉長期記憶": "找出值得長期保留的背景、規則、偏好。",
        "提煉使用者偏好": "聚焦使用者的習慣、偏好、風格。",
        "提煉專案背景": "聚焦專案、任務、長期脈絡。",
        "提煉工作規則": "聚焦可重複使用的規則、流程、判斷標準。",
    },
}

BUILTIN_WORKFLOW_TEMPLATES = {
    "文件快速整理": [
        {"mode": "Summary", "template": "重點摘要", "prompt": "請摘要目前上傳的檔案，先一句總結，再列 3 到 5 個重點，最後補充可執行建議。"},
        {"mode": "Report", "template": "正式報告", "prompt": "請根據目前這段對話與已有成果，整理成正式報告。"},
        {"mode": "Memory Assist", "template": "提煉長期記憶", "prompt": "請從目前這段對話與成果中，提煉值得長期保存的記憶。"},
    ],
    "表格檢查流程": [
        {"mode": "File Analyze", "template": "異常檢查", "prompt": "請檢查目前上傳表格中的異常、缺值、資料品質問題與風險。"},
        {"mode": "Chat", "template": "行動建議", "prompt": "請根據目前表格分析結果，產出清楚的下一步行動建議。"},
        {"mode": "Memory Assist", "template": "提煉工作規則", "prompt": "請從目前表格分析流程中，提煉值得保留的工作規則或判斷標準。"},
    ],
    "對話收斂流程": [
        {"mode": "Summary", "template": "條列整理", "prompt": "請整理目前這段對話，條列出重點與關鍵脈絡。"},
        {"mode": "Report", "template": "進度更新", "prompt": "請根據目前對話內容，整理成進度更新格式。"},
        {"mode": "Memory Assist", "template": "提煉專案背景", "prompt": "請從目前這段對話中，提煉專案背景與長期脈絡記憶。"},
    ],
}

FORMAT_TEMPLATES = {
    "主管版": "請把以下內容整理成主管版。要求：先結論、再列 3~5 個管理層最關心的重點、最後列決策建議。語氣精簡、務實、偏高層視角。",
    "客戶版": "請把以下內容整理成客戶版。要求：語氣完整、專業、易讀，避免內部黑話。輸出需可直接給客戶閱讀。",
    "簡報版": "請把以下內容整理成簡報版。要求：用投影片風格條列，分成數個段落或頁面，每頁 3~5 個 bullet，適合做簡報。",
    "條列版": "請把以下內容整理成條列版。要求：乾淨、精簡、快速閱讀，重點分層清楚。",
}

WORKFLOW_MEMORY_SOURCE = "custom_workflow_v1"
TASK_BOARD_SOURCE = "task_board_v2"
TASK_STATUS_OPTIONS = ["待處理", "進行中", "已完成", "已交付"]


# =========================================================
# Session State Defaults
# =========================================================
defaults = {
    "active_chat_id": None,
    "messages": [],
    "uploaded_file_cache": None,
    "memory_refresh_key": 0,
    "tool_panel": "Desk",
    "message_window": "最近 30 則",
    "work_mode": "Chat",
    "task_template": "一般問答",
    "last_memory_suggestion": None,
    "auto_memory_notice": None,
    "project_summary_cache": {},
    "workflow_queue": [],
    "workflow_name": None,
    "workflow_notice": None,
    "custom_workflows_cache": None,
    "workflow_refresh_key": 0,
    "task_board_cache": None,
    "task_board_refresh_key": 0,
    "daily_summary_cache": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# =========================================================
# Utility
# =========================================================
def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def parse_iso_dt(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except Exception:
        return None


def shorten_text(text: str, max_len: int = 80) -> str:
    t = (text or "").strip().replace("\n", " ")
    return t[:max_len] if t else "New Chat"


def tokenize(text: str) -> List[str]:
    text = (text or "").lower()
    tokens = re.findall(r"[\w\u4e00-\u9fff]+", text)
    stopwords = {
        "the", "a", "an", "to", "is", "are", "of", "and", "or", "in", "on",
        "我", "你", "他", "她", "它", "我們", "你們", "他們", "請", "幫我", "一下",
        "可以", "想", "要", "這個", "那個", "就是", "然後", "還有", "目前", "今天"
    }
    return [t for t in tokens if t not in stopwords and len(t) > 1]


def safe_json_loads(text: str) -> Optional[Dict[str, Any]]:
    text = text.strip()
    try:
        return json.loads(text)
    except Exception:
        pass

    code_match = re.search(r"```json\s*(\{.*?\})\s*```", text, re.S)
    if code_match:
        try:
            return json.loads(code_match.group(1))
        except Exception:
            return None

    brace_match = re.search(r"(\{.*\})", text, re.S)
    if brace_match:
        try:
            return json.loads(brace_match.group(1))
        except Exception:
            return None

    return None


def clip_preview(text: str, max_len: int = 220) -> str:
    t = re.sub(r"\s+", " ", text or "").strip()
    return t[:max_len] + "..." if len(t) > max_len else t


def parse_user_mode_template(content: str) -> Tuple[str, str, str]:
    match = re.match(r"^\[(.+?)\s*/\s*(.+?)\]\s*(.*)$", content or "", re.S)
    if match:
        return match.group(1).strip(), match.group(2).strip(), match.group(3).strip()
    return "Chat", "一般問答", content or ""


def get_task_options(mode: str) -> List[str]:
    options = list(TASK_TEMPLATES.get(mode, {}).keys())
    return options or ["一般問答"]


def ensure_valid_task_template():
    valid_options = get_task_options(st.session_state.work_mode)
    if st.session_state.task_template not in valid_options:
        st.session_state.task_template = valid_options[0]


def mode_help_text(mode: str) -> str:
    mapping = {
        "Chat": "一般聊天、提問、討論、文字協助。",
        "Summary": "把長文、對話、檔案整理成精簡重點。",
        "File Analyze": "分析 Excel / CSV / PDF / txt 內容與重點。",
        "Report": "輸出成較正式、可交付的報告格式。",
        "Memory Assist": "把內容整理成可長期記住的記憶建議。",
    }
    return mapping.get(mode, "")


def task_help_text(mode: str, task_template: str) -> str:
    return TASK_TEMPLATES.get(mode, {}).get(task_template, "")


def get_window_size(option: str, total_count: int) -> int:
    mapping = {
        "最近 20 則": 20,
        "最近 30 則": 30,
        "最近 50 則": 50,
        "全部顯示": total_count,
    }
    return mapping.get(option, 30)


def all_mode_options() -> List[str]:
    return ["Chat", "Summary", "File Analyze", "Report", "Memory Assist"]


def sanitize_filename(text: str, max_len: int = 40) -> str:
    s = re.sub(r'[\\/:*?"<>|]+', "_", (text or "").strip())
    s = re.sub(r"\s+", "_", s)
    s = s[:max_len].strip("_")
    return s or "lobster_output"


def get_today_start_utc() -> datetime:
    now = now_utc()
    return now.replace(hour=0, minute=0, second=0, microsecond=0)


def dedupe_linked_outputs(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    result = []
    for item in items:
        key = (item.get("output_key", ""), item.get("label", ""))
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


# =========================================================
# DB Helpers
# =========================================================
def list_sessions() -> List[Dict]:
    resp = (
        supabase.table("lobster_sessions")
        .select("id, title, created_at, updated_at")
        .order("updated_at", desc=True)
        .limit(200)
        .execute()
    )
    return resp.data or []


def create_session(title: str = "New Chat") -> str:
    sid = str(uuid.uuid4())
    supabase.table("lobster_sessions").insert({"id": sid, "title": title[:80] or "New Chat"}).execute()
    return sid


def create_message(session_id: str, role: str, content: str, status: str = "completed", error_message: Optional[str] = None) -> Optional[int]:
    resp = (
        supabase.table("lobster_messages")
        .insert(
            {
                "session_id": session_id,
                "role": role,
                "content": content,
                "status": status,
                "error_message": error_message,
            }
        )
        .execute()
    )
    if resp.data and len(resp.data) > 0:
        return resp.data[0]["id"]
    return None


def update_message(message_id: int, content: str, status: str, error_message: Optional[str] = None) -> None:
    (
        supabase.table("lobster_messages")
        .update({"content": content, "status": status, "error_message": error_message})
        .eq("id", message_id)
        .execute()
    )


def mark_stale_pending_messages(session_id: str, stale_minutes: int = 5) -> None:
    resp = (
        supabase.table("lobster_messages")
        .select("id, status, updated_at")
        .eq("session_id", session_id)
        .eq("role", "assistant")
        .eq("status", "pending")
        .execute()
    )
    rows = resp.data or []
    cutoff = now_utc() - timedelta(minutes=stale_minutes)
    for row in rows:
        updated_at = parse_iso_dt(row.get("updated_at"))
        if updated_at and updated_at < cutoff:
            update_message(
                row["id"],
                content="⚠️ 上一次回覆在處理途中中斷，請重新發問。",
                status="failed",
                error_message="interrupted_or_closed",
            )


def load_messages(session_id: str) -> None:
    mark_stale_pending_messages(session_id)
    resp = (
        supabase.table("lobster_messages")
        .select("id, role, content, status, error_message, created_at, updated_at")
        .eq("session_id", session_id)
        .order("created_at")
        .execute()
    )
    st.session_state.messages = [
        {
            "id": row["id"],
            "role": row["role"],
            "content": row["content"],
            "status": row.get("status", "completed"),
            "error_message": row.get("error_message"),
            "created_at": row.get("created_at"),
            "updated_at": row.get("updated_at"),
        }
        for row in (resp.data or [])
    ]


def update_session_title_if_needed(session_id: str, prompt: str) -> None:
    resp = supabase.table("lobster_sessions").select("title").eq("id", session_id).single().execute()
    current_title = "New Chat"
    if resp.data:
        current_title = resp.data.get("title") or "New Chat"
    if current_title == "New Chat":
        supabase.table("lobster_sessions").update({"title": shorten_text(prompt, 30)}).eq("id", session_id).execute()


def rename_session(session_id: str, new_title: str) -> None:
    supabase.table("lobster_sessions").update({"title": shorten_text(new_title, 80)}).eq("id", session_id).execute()


def delete_session(session_id: str) -> None:
    supabase.table("lobster_sessions").delete().eq("id", session_id).execute()


def touch_session(session_id: str) -> None:
    resp = supabase.table("lobster_sessions").select("title").eq("id", session_id).single().execute()
    if resp.data:
        current_title = resp.data.get("title") or "New Chat"
        supabase.table("lobster_sessions").update({"title": current_title}).eq("id", session_id).execute()


def ensure_active_chat() -> None:
    if st.session_state.active_chat_id:
        return
    sessions = list_sessions()
    if sessions:
        st.session_state.active_chat_id = sessions[0]["id"]
        load_messages(st.session_state.active_chat_id)
    else:
        sid = create_session("New Chat")
        st.session_state.active_chat_id = sid
        st.session_state.messages = []


# =========================================================
# Memory Helpers
# =========================================================
def list_memory_items(limit: int = 100) -> List[Dict]:
    resp = (
        supabase.table("memory_items")
        .select("id, title, content, tags, source, created_at, updated_at")
        .order("updated_at", desc=True)
        .limit(limit)
        .execute()
    )
    return resp.data or []


def create_memory_item(title: str, content: str, tags: Optional[List[str]] = None, source: str = "manual") -> None:
    clean_tags = [t.strip() for t in (tags or []) if str(t).strip()]
    supabase.table("memory_items").insert(
        {"title": shorten_text(title, 120), "content": content.strip(), "tags": clean_tags, "source": source}
    ).execute()


def update_memory_item(memory_id: str, title: str, content: str, tags: Optional[List[str]] = None, source: Optional[str] = None) -> None:
    payload = {
        "title": shorten_text(title, 120),
        "content": content.strip(),
        "tags": [t.strip() for t in (tags or []) if str(t).strip()],
    }
    if source is not None:
        payload["source"] = source
    supabase.table("memory_items").update(payload).eq("id", memory_id).execute()


def delete_memory_item(memory_id: str) -> None:
    supabase.table("memory_items").delete().eq("id", memory_id).execute()


def get_relevant_memory(query: str, top_k: int = 5) -> List[Dict]:
    memories = list_memory_items(limit=200)
    if not memories:
        return []
    query_tokens = set(tokenize(query))
    scored = []
    for item in memories:
        body = f"{item.get('title', '')}\n{item.get('content', '')}\n{' '.join(item.get('tags', []) or [])}"
        body_tokens = set(tokenize(body))
        overlap = len(query_tokens.intersection(body_tokens))
        bonus = 5 if query and query.lower() in body.lower() else 0
        score = overlap + bonus
        if score > 0:
            scored.append((score, item))
    scored.sort(key=lambda x: (x[0], x[1].get("updated_at", "")), reverse=True)
    return [item for _, item in scored[:top_k]]


def format_memory_context(query: str) -> str:
    relevant = get_relevant_memory(query, top_k=5)
    if not relevant:
        return ""
    parts = ["以下是與使用者問題最相關的長期記憶，回答時可優先參考："]
    for idx, item in enumerate(relevant, start=1):
        parts.append(
            f"""記憶 {idx}
標題：{item.get('title', '')}
標籤：{', '.join(item.get('tags', []) or [])}
內容：
{item.get('content', '')}"""
        )
    return "\n\n".join(parts)


# =========================================================
# Custom Workflow Storage
# =========================================================
def normalize_workflow_payload(payload: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return None
    name = str(payload.get("name", "")).strip()
    steps = payload.get("steps", [])
    if not name or not isinstance(steps, list) or not steps:
        return None

    clean_steps = []
    for step in steps[:5]:
        if not isinstance(step, dict):
            continue
        mode = str(step.get("mode", "")).strip()
        template = str(step.get("template", "")).strip()
        prompt = str(step.get("prompt", "")).strip()
        if mode and template and prompt:
            clean_steps.append({"mode": mode, "template": template, "prompt": prompt})

    if not clean_steps:
        return None

    return {"name": name, "steps": clean_steps}


def load_custom_workflows_from_memory(force_refresh: bool = False) -> Dict[str, List[Dict[str, str]]]:
    if st.session_state.custom_workflows_cache is not None and not force_refresh:
        return st.session_state.custom_workflows_cache

    memories = list_memory_items(limit=500)
    workflows: Dict[str, List[Dict[str, str]]] = {}

    for item in memories:
        if item.get("source") != WORKFLOW_MEMORY_SOURCE:
            continue
        payload = safe_json_loads(item.get("content", "") or "")
        normalized = normalize_workflow_payload(payload)
        if normalized:
            workflows[normalized["name"]] = normalized["steps"]

    st.session_state.custom_workflows_cache = workflows
    return workflows


def save_custom_workflow(name: str, steps: List[Dict[str, str]]) -> None:
    payload = {"name": name.strip(), "steps": steps}
    create_memory_item(
        title=f"Workflow｜{name.strip()}",
        content=json.dumps(payload, ensure_ascii=False),
        tags=["workflow", "custom"],
        source=WORKFLOW_MEMORY_SOURCE,
    )
    st.session_state.custom_workflows_cache = None


def delete_custom_workflow(name: str) -> None:
    memories = list_memory_items(limit=500)
    for item in memories:
        if item.get("source") != WORKFLOW_MEMORY_SOURCE:
            continue
        payload = safe_json_loads(item.get("content", "") or "")
        normalized = normalize_workflow_payload(payload)
        if normalized and normalized["name"] == name:
            delete_memory_item(item["id"])
    st.session_state.custom_workflows_cache = None


# =========================================================
# Task Board Storage
# =========================================================
def normalize_linked_output_payload(item: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(item, dict):
        return None
    label = str(item.get("label", "")).strip()
    preview = str(item.get("preview", "")).strip()
    content = str(item.get("content", "")).strip()
    output_key = str(item.get("output_key", "")).strip()
    mode = str(item.get("mode", "")).strip()
    template = str(item.get("template", "")).strip()
    if not label:
        label = "未命名成果"
    if not output_key:
        output_key = label
    return {
        "label": label,
        "preview": preview,
        "content": content,
        "output_key": output_key,
        "mode": mode,
        "template": template,
    }


def normalize_task_payload(payload: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return None
    title = str(payload.get("title", "")).strip()
    description = str(payload.get("description", "")).strip()
    status = str(payload.get("status", "")).strip()
    linked_outputs_raw = payload.get("linked_outputs", [])
    if not title:
        return None
    if status not in TASK_STATUS_OPTIONS:
        status = "待處理"

    linked_outputs = []
    if isinstance(linked_outputs_raw, list):
        for item in linked_outputs_raw:
            normalized_item = normalize_linked_output_payload(item)
            if normalized_item:
                linked_outputs.append(normalized_item)

    linked_outputs = dedupe_linked_outputs(linked_outputs)

    return {
        "title": title,
        "description": description,
        "status": status,
        "linked_outputs": linked_outputs,
    }


def load_task_board(force_refresh: bool = False) -> List[Dict[str, Any]]:
    if st.session_state.task_board_cache is not None and not force_refresh:
        return st.session_state.task_board_cache

    memories = list_memory_items(limit=500)
    tasks = []
    for item in memories:
        if item.get("source") != TASK_BOARD_SOURCE:
            continue
        payload = safe_json_loads(item.get("content", "") or "")
        normalized = normalize_task_payload(payload)
        if normalized:
            tasks.append(
                {
                    "memory_id": item["id"],
                    "title": normalized["title"],
                    "description": normalized["description"],
                    "status": normalized["status"],
                    "linked_outputs": normalized["linked_outputs"],
                    "created_at": item.get("created_at", ""),
                    "updated_at": item.get("updated_at", ""),
                }
            )

    tasks.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    st.session_state.task_board_cache = tasks
    return tasks


def build_task_payload(title: str, description: str, status: str, linked_outputs: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    return {
        "title": title.strip(),
        "description": description.strip(),
        "status": status if status in TASK_STATUS_OPTIONS else "待處理",
        "linked_outputs": dedupe_linked_outputs(linked_outputs or []),
    }


def create_task_item(title: str, description: str = "", status: str = "待處理", linked_outputs: Optional[List[Dict[str, Any]]] = None) -> None:
    payload = build_task_payload(title, description, status, linked_outputs)
    create_memory_item(
        title=f"Task｜{title.strip()}",
        content=json.dumps(payload, ensure_ascii=False),
        tags=["task", payload["status"]],
        source=TASK_BOARD_SOURCE,
    )
    st.session_state.task_board_cache = None


def update_task_item(memory_id: str, title: str, description: str, status: str, linked_outputs: Optional[List[Dict[str, Any]]] = None) -> None:
    payload = build_task_payload(title, description, status, linked_outputs)
    update_memory_item(
        memory_id=memory_id,
        title=f"Task｜{title.strip()}",
        content=json.dumps(payload, ensure_ascii=False),
        tags=["task", payload["status"]],
        source=TASK_BOARD_SOURCE,
    )
    st.session_state.task_board_cache = None


def delete_task_item(memory_id: str) -> None:
    delete_memory_item(memory_id)
    st.session_state.task_board_cache = None


def build_output_link_record(output: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "label": f"{output['mode']} / {output['template']}",
        "preview": output["assistant_preview"],
        "content": output["assistant_full"],
        "output_key": str(output["id"]),
        "mode": output["mode"],
        "template": output["template"],
    }


def add_output_to_existing_task(task_memory_id: str, output: Dict[str, Any]) -> None:
    tasks = load_task_board(force_refresh=True)
    target = next((t for t in tasks if t["memory_id"] == task_memory_id), None)
    if not target:
        return
    new_links = target["linked_outputs"] + [build_output_link_record(output)]
    new_links = dedupe_linked_outputs(new_links)
    update_task_item(
        memory_id=target["memory_id"],
        title=target["title"],
        description=target["description"],
        status=target["status"],
        linked_outputs=new_links,
    )


# =========================================================
# File Helpers
# =========================================================
def dataframe_preview_markdown(df: pd.DataFrame, max_rows: int = 8, max_cols: int = 12) -> str:
    clipped = df.iloc[:max_rows, :max_cols].copy().fillna("")
    try:
        return clipped.to_markdown(index=False)
    except Exception:
        return clipped.to_csv(index=False)


def dataframe_profile_text(df: pd.DataFrame, name: str) -> str:
    parts = [f"=== 資料表：{name} ===", f"列數：{len(df)}", f"欄數：{len(df.columns)}"]
    if len(df.columns) > 0:
        parts.append("欄位名稱：")
        parts.append(", ".join([str(c) for c in df.columns[:50]]))

        dtype_lines = [f"- {col}: {str(df[col].dtype)}" for col in df.columns[:30]]
        if dtype_lines:
            parts.append("欄位型別：")
            parts.extend(dtype_lines)

        null_counts = df.isna().sum()
        null_lines = [f"- {col}: 缺值 {int(null_counts[col])}" for col in df.columns[:30] if int(null_counts[col]) > 0]
        if null_lines:
            parts.append("缺值統計：")
            parts.extend(null_lines)
        else:
            parts.append("缺值統計：未發現明顯缺值")

        numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
        if numeric_cols:
            parts.append("數值欄摘要：")
            for col in numeric_cols[:10]:
                series = pd.to_numeric(df[col], errors="coerce")
                mean_val = round(series.mean(), 4) if pd.notna(series.mean()) else "NA"
                parts.append(f"- {col}: min={series.min()}, max={series.max()}, mean={mean_val}")
        else:
            parts.append("數值欄摘要：無數值欄")

        parts.append("前幾列預覽：")
        parts.append(dataframe_preview_markdown(df))
    return "\n".join(parts)


def excel_bytes_to_text(file_bytes: bytes, filename: str) -> str:
    output_parts = [f"Excel file: {filename}"]
    xls = pd.ExcelFile(BytesIO(file_bytes))
    for sheet_name in xls.sheet_names[:5]:
        try:
            df = pd.read_excel(BytesIO(file_bytes), sheet_name=sheet_name)
            output_parts.append(dataframe_profile_text(df, sheet_name))
        except Exception as e:
            output_parts.append(f"=== 資料表：{sheet_name} ===\n讀取失敗：{e}")
    return "\n\n".join(output_parts)


def csv_bytes_to_text(file_bytes: bytes, filename: str) -> str:
    try:
        df = pd.read_csv(BytesIO(file_bytes))
        return f"CSV file: {filename}\n\n{dataframe_profile_text(df, filename)}"
    except Exception as e:
        return f"CSV file: {filename}\n讀取失敗：{e}"


def text_bytes_to_text(file_bytes: bytes, filename: str) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        return f"Text file: {filename}\n\n{text[:30000]}"
    except Exception as e:
        return f"Text file: {filename}\n讀取失敗：{e}"


def build_file_context_parts() -> Tuple[List, str]:
    cached = st.session_state.uploaded_file_cache
    if not cached:
        return [], ""

    filename = cached["name"]
    mime_type = cached["type"]
    file_bytes = cached["data"]
    ext = filename.lower().split(".")[-1] if "." in filename else ""

    content_parts: List[Any] = []
    file_text_context = ""

    if ext in ["xlsx", "xls"]:
        file_text_context = "以下是使用者上傳 Excel 的結構化分析：\n" + excel_bytes_to_text(file_bytes, filename)
    elif ext == "csv":
        file_text_context = "以下是使用者上傳 CSV 的結構化分析：\n" + csv_bytes_to_text(file_bytes, filename)
    elif ext == "txt":
        file_text_context = "以下是使用者上傳文字檔內容：\n" + text_bytes_to_text(file_bytes, filename)
    elif ext in ["pdf", "png", "jpg", "jpeg"]:
        content_parts.append({"mime_type": mime_type, "data": file_bytes})
        file_text_context = f"使用者已上傳檔案：{filename}。請結合多模態內容理解。"
    else:
        file_text_context = f"使用者上傳了一個檔案：{filename}，但目前系統無法完整解析這個格式。"

    return content_parts, file_text_context


# =========================================================
# AI Prompt Helpers
# =========================================================
BASE_ASSISTANT_RULES = """
你是使用者的私人 AI 助手「龍蝦王助手」。
請永遠使用繁體中文回答。
請先講結論，再補充細節。
不要空話，不要過度客套。
如果不確定，請明確說不確定。
""".strip()


def get_mode_instruction(mode: str) -> str:
    prompts = {
        "Chat": "你現在是一般聊天助理模式。先結論，再補充，簡潔清楚。",
        "Summary": "你現在是摘要模式。先一句話總結，再列 3 到 5 個重點，最後給可執行建議。",
        "File Analyze": "你現在是檔案分析模式。先概覽，再列結構重點，再指出異常與風險，最後給下一步。",
        "Report": "你現在是報告模式。輸出結構：標題、背景、重點整理、結論、建議/下一步。",
        "Memory Assist": "你現在是記憶助理模式。請判斷哪些資訊值得長期保存成記憶。",
    }
    return prompts.get(mode, prompts["Chat"])


def build_content_parts(prompt: str, mode: str, task_template: str) -> List:
    content_parts: List[Any] = []
    memory_context = format_memory_context(prompt)
    file_parts, file_text_context = build_file_context_parts()
    task_instruction = TASK_TEMPLATES.get(mode, {}).get(task_template, "")

    system_block = (
        BASE_ASSISTANT_RULES
        + "\n\n"
        + get_mode_instruction(mode)
        + "\n\n"
        + f"目前任務模板：{task_template}\n任務模板要求：{task_instruction}"
    )
    content_parts.append(system_block)

    if memory_context:
        content_parts.append("\n\n" + memory_context)
    if file_text_context:
        content_parts.append("\n\n" + file_text_context)

    content_parts.append("\n\n使用者要求如下：\n" + prompt)

    if file_parts:
        content_parts.extend(file_parts)

    return content_parts


def build_history_for_model(messages: List[Dict], limit: int = 12) -> List[Dict]:
    history = []
    filtered = [m for m in messages if m["role"] in ["user", "assistant"]]
    for m in filtered[-limit:]:
        content = m.get("content", "")
        if not content:
            continue
        role = "user" if m["role"] == "user" else "model"
        history.append({"role": role, "parts": [content]})
    return history


def get_candidate_models() -> List[str]:
    return ["gemini-2.5-flash-lite", "gemini-2.5-flash"]


def is_quota_error(err_text: str) -> bool:
    t = err_text.lower()
    return (
        "429" in t
        or "quota" in t
        or "rate limit" in t
        or "resource has been exhausted" in t
        or "exceeded your current quota" in t
    )


def is_model_not_found(err_text: str) -> bool:
    t = err_text.lower()
    return "404" in t or "not found" in t or "unsupported" in t or "is not supported" in t


def generate_reply_with_key_rotation(history: List[Dict], content_parts: List, preferred_models: Optional[List[str]] = None) -> Tuple[str, Optional[str], Optional[str]]:
    last_error = None
    model_candidates = preferred_models or get_candidate_models()

    for key_idx, api_key in enumerate(API_KEY_LIST, start=1):
        genai.configure(api_key=api_key)
        for model_name in model_candidates:
            try:
                model = genai.GenerativeModel(model_name)
                for attempt in range(2):
                    try:
                        response = model.generate_content(contents=history + [{"role": "user", "parts": content_parts}])
                        reply = getattr(response, "text", None)
                        if not reply:
                            reply = "⚠️ 龍蝦有收到問題，但這次模型沒有回傳文字內容。"
                        return reply, model_name, f"key-{key_idx}"
                    except Exception as e:
                        err_text = str(e)
                        last_error = err_text
                        if is_quota_error(err_text) and attempt == 0:
                            time.sleep(1.2)
                            continue
                        raise
            except Exception as e:
                err_text = str(e)
                last_error = err_text
                if is_quota_error(err_text) or is_model_not_found(err_text):
                    continue
                continue

    if last_error and is_quota_error(last_error):
        return "🦞 今天 AI 額度暫時用完了，請稍後再試，或新增更多 API key 輪替。", None, None
    return f"⚠️ AI error: {last_error}" if last_error else "⚠️ AI 暫時無法回覆，請稍後再試。", None, None


# =========================================================
# Summary / Outputs Helpers
# =========================================================
def generate_project_summary(messages: List[Dict]) -> Optional[Dict[str, Any]]:
    if not messages:
        return None

    text_chunks = []
    for m in messages[-24:]:
        role = "使用者" if m["role"] == "user" else "助手"
        content = m.get("content", "")
        if content:
            text_chunks.append(f"{role}：{content}")

    convo = "\n\n".join(text_chunks)[:22000]

    prompt = f"""
請根據以下對話，整理成專案摘要。
只輸出 JSON，不要加任何其他文字。

JSON 格式：
{{
  "project_name": "一句短標題",
  "current_focus": "目前主要在處理什麼",
  "completed": ["已完成1", "已完成2"],
  "missing": ["還缺1", "還缺2"],
  "next_steps": ["下一步1", "下一步2"],
  "status_note": "一句整體狀態"
}}

對話內容：
{convo}
""".strip()

    reply, _, _ = generate_reply_with_key_rotation(
        history=[],
        content_parts=[prompt],
        preferred_models=["gemini-2.5-flash-lite", "gemini-2.5-flash"],
    )

    if reply.startswith("⚠️") or reply.startswith("🦞 今天 AI 額度"):
        return None

    return safe_json_loads(reply)


def get_chat_summary_cache_key() -> str:
    return f"{st.session_state.active_chat_id}:{len(st.session_state.messages)}"


def get_or_refresh_project_summary(force_refresh: bool = False) -> Optional[Dict[str, Any]]:
    cache_key = get_chat_summary_cache_key()
    cache = st.session_state.project_summary_cache

    if not force_refresh and cache_key in cache:
        return cache[cache_key]

    summary = generate_project_summary(st.session_state.messages)
    if summary:
        cache[cache_key] = summary
        st.session_state.project_summary_cache = cache
    return summary


def get_output_records(messages: List[Dict]) -> List[Dict]:
    records = []
    prev_user = None
    record_id = 0
    for msg in messages:
        if msg["role"] == "user":
            prev_user = msg["content"]
        elif msg["role"] == "assistant" and prev_user and msg.get("status", "completed") == "completed":
            mode, template, user_prompt = parse_user_mode_template(prev_user)
            record_id += 1
            records.append(
                {
                    "id": f"out_{record_id}",
                    "mode": mode,
                    "template": template,
                    "user_prompt": user_prompt,
                    "assistant_full": msg["content"],
                    "assistant_preview": clip_preview(msg["content"], 260),
                    "created_at": msg.get("created_at", ""),
                }
            )
            prev_user = None
    return list(reversed(records))


def get_recent_output_cards(messages: List[Dict], max_cards: int = 6) -> List[Dict]:
    return get_output_records(messages)[:max_cards]


def get_suggested_next_actions(messages: List[Dict], has_file: bool) -> List[str]:
    suggestions = []
    if has_file:
        suggestions.append("可直接摘要、分析或檢查目前檔案。")
        suggestions.append("可把目前檔案轉成正式報告，方便交付。")
    if messages:
        suggestions.append("可將目前對話或成果提煉成長期記憶。")
        suggestions.append("可針對目前內容產出下一步行動建議。")
    if not messages and not has_file:
        suggestions.append("先上傳一份檔案，或直接輸入你想處理的任務。")
        suggestions.append("先選模式與任務模板，讓龍蝦用正確方式工作。")
    return suggestions[:4]


# =========================================================
# Daily Summary Helpers
# =========================================================
def is_today(dt_str: Optional[str]) -> bool:
    dt = parse_iso_dt(dt_str)
    if not dt:
        return False
    return dt >= get_today_start_utc()


def build_daily_summary(force_refresh: bool = False) -> Dict[str, Any]:
    if st.session_state.daily_summary_cache is not None and not force_refresh:
        return st.session_state.daily_summary_cache

    messages = st.session_state.messages
    outputs = get_output_records(messages)
    tasks = load_task_board(force_refresh=False)

    todays_outputs = [o for o in outputs if is_today(o.get("created_at"))]
    todays_new_tasks = [t for t in tasks if is_today(t.get("created_at"))]
    todays_updated_tasks = [t for t in tasks if is_today(t.get("updated_at"))]

    done_tasks = [t for t in todays_updated_tasks if t["status"] in ["已完成", "已交付"]]
    blocked_tasks = [t for t in tasks if t["status"] in ["待處理", "進行中"]][:6]

    completed_titles = [shorten_text(t["title"], 50) for t in done_tasks[:6]]
    new_task_titles = [shorten_text(t["title"], 50) for t in todays_new_tasks[:6]]
    output_titles = [f"{o['mode']} / {o['template']}" for o in todays_outputs[:6]]
    blocked_titles = [shorten_text(t["title"], 50) for t in blocked_tasks[:6]]

    suggestion_pool = []
    if blocked_tasks:
        suggestion_pool.append("優先處理看板中『進行中』但尚未完成的任務。")
    if todays_outputs and not done_tasks:
        suggestion_pool.append("今天已有成果產出，可考慮掛到任務並推進成已完成。")
    if not todays_outputs and not todays_new_tasks:
        suggestion_pool.append("今天還沒有明顯工作紀錄，可先新增任務或產出第一份成果。")
    if todays_outputs:
        suggestion_pool.append("可將今日成果格式化後下載，準備對內或對外交付。")
    if len(tasks) > 0:
        suggestion_pool.append("檢查待處理項目，挑 1 到 2 件最重要的先收尾。")

    summary = {
        "date_label": now_utc().strftime("%Y-%m-%d"),
        "new_tasks_count": len(todays_new_tasks),
        "done_tasks_count": len(done_tasks),
        "outputs_count": len(todays_outputs),
        "blocked_count": len(blocked_tasks),
        "new_tasks": new_task_titles,
        "done_tasks": completed_titles,
        "outputs": output_titles,
        "blocked": blocked_titles,
        "next_steps": suggestion_pool[:4],
    }

    st.session_state.daily_summary_cache = summary
    return summary


# =========================================================
# Memory Extraction Helpers
# =========================================================
def extract_memory_suggestion(reply: str) -> Optional[Dict]:
    if "是否值得記住" not in reply:
        return None

    def pick(pattern: str) -> str:
        m = re.search(pattern, reply, re.S)
        return m.group(1).strip() if m else ""

    worth = pick(r"是否值得記住[:：]\s*(.+?)(?:\n|$)")
    title = pick(r"建議標題[:：]\s*(.+?)(?:\n|$)")
    tags_line = pick(r"建議標籤[:：]\s*(.+?)(?:\n|$)")
    content = pick(r"建議記憶內容[:：]\s*(.+?)(?:\n\s*\d+\.|$)")
    reason = pick(r"理由[:：]\s*(.+?)(?:\n|$)")
    tags = [t.strip() for t in re.split(r"[，,、]", tags_line) if t.strip()]

    if not any([worth, title, content]):
        return None

    return {
        "worth": worth,
        "title": title or "未命名記憶",
        "tags": tags,
        "content": content,
        "reason": reason,
    }


def auto_extract_memory_from_conversation(user_prompt: str, assistant_reply: str, mode: str, task_template: str) -> Optional[Dict]:
    auto_prompt = f"""
你現在要做的是：判斷以下內容是否值得存成「長期記憶」。

規則：
1. 只保留會在未來仍有價值的資訊
2. 優先保留：偏好、長期專案、工作方式、固定規則、重要背景
3. 不要保留短期閒聊、一次性任務、臨時問答
4. 僅輸出 JSON，不要加解釋

JSON 格式：
{{
  "should_save": true 或 false,
  "title": "記憶標題",
  "tags": ["tag1", "tag2"],
  "content": "記憶內容",
  "reason": "為何值得或不值得保存"
}}

模式：{mode}
任務模板：{task_template}

使用者：
{user_prompt}

助手回覆：
{assistant_reply}
""".strip()

    reply, _, _ = generate_reply_with_key_rotation(
        history=[],
        content_parts=[auto_prompt],
        preferred_models=["gemini-2.5-flash-lite", "gemini-2.5-flash"],
    )

    if reply.startswith("⚠️") or reply.startswith("🦞 今天 AI 額度"):
        return None

    parsed = safe_json_loads(reply)
    if not parsed:
        return None

    should_save = bool(parsed.get("should_save", False))
    title = str(parsed.get("title", "")).strip()
    content = str(parsed.get("content", "")).strip()
    tags = parsed.get("tags", [])
    if not isinstance(tags, list):
        tags = []

    if not should_save or not content:
        return None

    return {
        "title": title or content[:30],
        "content": content,
        "tags": [str(t).strip() for t in tags if str(t).strip()],
        "reason": str(parsed.get("reason", "")).strip(),
    }


# =========================================================
# Format Helpers
# =========================================================
def format_output_prompt(source_text: str, format_name: str) -> str:
    instruction = FORMAT_TEMPLATES.get(format_name, "")
    return f"{instruction}\n\n以下是原始內容：\n\n{source_text}"


# =========================================================
# Download Helpers
# =========================================================
def build_txt_bytes(text: str) -> bytes:
    return (text or "").encode("utf-8")


def build_md_bytes(text: str) -> bytes:
    return (text or "").encode("utf-8")


def build_docx_bytes(title: str, text: str) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for block in (text or "").split("\n\n"):
        doc.add_paragraph(block)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_pdf_bytes(title: str, text: str) -> bytes:
    bio = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(bio, pagesize=A4)
    _, height = A4
    x = 40
    y = height - 50
    line_height = 16

    c.setFont("STSong-Light", 14)
    if title:
        c.drawString(x, y, title)
        y -= 24

    c.setFont("STSong-Light", 11)
    raw_lines = []
    for para in (text or "").split("\n"):
        if para.strip() == "":
            raw_lines.append("")
            continue
        while len(para) > 48:
            raw_lines.append(para[:48])
            para = para[48:]
        raw_lines.append(para)

    for line in raw_lines:
        if y < 40:
            c.showPage()
            c.setFont("STSong-Light", 11)
            y = height - 40
        c.drawString(x, y, line)
        y -= line_height

    c.save()
    bio.seek(0)
    return bio.getvalue()


def get_download_payload(text: str, base_name: str) -> Dict[str, Dict[str, Any]]:
    safe_name = sanitize_filename(base_name)
    return {
        "txt": {"label": "下載 TXT", "data": build_txt_bytes(text), "file_name": f"{safe_name}.txt", "mime": "text/plain"},
        "md": {"label": "下載 MD", "data": build_md_bytes(text), "file_name": f"{safe_name}.md", "mime": "text/markdown"},
        "docx": {
            "label": "下載 DOCX",
            "data": build_docx_bytes(base_name, text),
            "file_name": f"{safe_name}.docx",
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        "pdf": {"label": "下載 PDF", "data": build_pdf_bytes(base_name, text), "file_name": f"{safe_name}.pdf", "mime": "application/pdf"},
    }


# =========================================================
# Init
# =========================================================
ensure_active_chat()
ensure_valid_task_template()


# =========================================================
# Core Runner
# =========================================================
def run_assistant_turn(prompt: str, mode: str, task_template: str, should_rerun: bool = True) -> None:
    session_id = st.session_state.active_chat_id
    st.session_state.auto_memory_notice = None
    st.session_state.daily_summary_cache = None

    user_display = f"[{mode} / {task_template}] {prompt}"
    st.session_state.messages.append({"role": "user", "content": user_display, "status": "completed", "error_message": None})

    try:
        create_message(session_id, "user", user_display, status="completed")
        update_session_title_if_needed(session_id, prompt)
        touch_session(session_id)
    except Exception:
        pass

    pending_text = "🦞 龍蝦正在思考中..."
    pending_id = None
    try:
        pending_id = create_message(session_id, "assistant", pending_text, status="pending")
    except Exception:
        pass

    st.session_state.messages.append({"id": pending_id, "role": "assistant", "content": pending_text, "status": "pending", "error_message": None})

    reply = ""
    final_status = "failed"
    final_error = None

    try:
        content_parts = build_content_parts(prompt, mode, task_template)
        history = build_history_for_model(st.session_state.messages[:-1], limit=12)
        reply, _, _ = generate_reply_with_key_rotation(history, content_parts)

        final_status = "failed" if reply.startswith("⚠️") or reply.startswith("🦞 今天 AI 額度") else "completed"
        final_error = reply if final_status == "failed" else None

        if final_status == "completed" and mode == "Memory Assist":
            st.session_state.last_memory_suggestion = extract_memory_suggestion(reply)

    except Exception as e:
        reply = f"⚠️ AI error: {e}"
        final_status = "failed"
        final_error = str(e)

    st.session_state.messages[-1] = {
        "id": pending_id,
        "role": "assistant",
        "content": reply,
        "status": final_status,
        "error_message": final_error,
    }

    try:
        if pending_id is not None:
            update_message(pending_id, content=reply, status=final_status, error_message=final_error)
        else:
            create_message(session_id, "assistant", reply, status=final_status, error_message=final_error)
        touch_session(session_id)
    except Exception:
        pass

    if final_status == "completed" and mode != "Memory Assist":
        try:
            auto_mem = auto_extract_memory_from_conversation(prompt, reply, mode, task_template)
            if auto_mem:
                create_memory_item(
                    title=auto_mem["title"],
                    content=auto_mem["content"],
                    tags=auto_mem["tags"],
                    source="auto",
                )
                st.session_state.auto_memory_notice = auto_mem["title"]
        except Exception:
            pass

    if should_rerun:
        st.rerun()


# =========================================================
# Workflow Engine
# =========================================================
def get_all_workflows() -> Dict[str, List[Dict[str, str]]]:
    workflows = dict(BUILTIN_WORKFLOW_TEMPLATES)
    custom = load_custom_workflows_from_memory(force_refresh=False)
    workflows.update(custom)
    return workflows


def start_workflow(workflow_name: str) -> bool:
    workflows = get_all_workflows()
    steps = workflows.get(workflow_name, [])
    if not steps:
        return False

    if workflow_name in ["文件快速整理"] and not st.session_state.uploaded_file_cache:
        st.warning("這個流程需要先上傳檔案。")
        return False

    if workflow_name in ["表格檢查流程"]:
        if not st.session_state.uploaded_file_cache:
            st.warning("這個流程需要先上傳 Excel 或 CSV。")
            return False
        fname = st.session_state.uploaded_file_cache["name"].lower()
        if not (fname.endswith(".csv") or fname.endswith(".xlsx") or fname.endswith(".xls")):
            st.warning("這個流程只適用 CSV / Excel。")
            return False

    if workflow_name in ["對話收斂流程"] and not st.session_state.messages:
        st.warning("這個流程需要先有一些對話內容。")
        return False

    st.session_state.workflow_queue = steps.copy()
    st.session_state.workflow_name = workflow_name
    st.session_state.workflow_notice = f"流程已啟動：{workflow_name}"
    st.rerun()
    return True


def process_workflow_step_if_needed():
    queue = st.session_state.workflow_queue
    if not queue:
        return

    step = queue.pop(0)
    st.session_state.workflow_queue = queue
    st.session_state.work_mode = step["mode"]
    st.session_state.task_template = step["template"]

    if not queue:
        st.session_state.workflow_notice = f"流程完成：{st.session_state.workflow_name}"
        st.session_state.workflow_name = None
    else:
        st.session_state.workflow_notice = f"流程進行中：{st.session_state.workflow_name}（剩餘 {len(queue)} 步）"

    run_assistant_turn(prompt=step["prompt"], mode=step["mode"], task_template=step["template"], should_rerun=True)


if st.session_state.workflow_queue:
    process_workflow_step_if_needed()


# =========================================================
# Header
# =========================================================
header_left, header_right = st.columns([5.2, 2.8])

with header_left:
    st.markdown(
        """
        <div class="lobster-header">
            <div class="lobster-brand">🦞 龍蝦王助手</div>
            <div class="lobster-subtitle">私人 AI 助手 · 任務與成果連動 V2.3</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with header_right:
    if st.button("＋ 新對話", use_container_width=True):
        sid = create_session("New Chat")
        st.session_state.active_chat_id = sid
        st.session_state.messages = []
        st.session_state.uploaded_file_cache = None
        st.session_state.last_memory_suggestion = None
        st.session_state.auto_memory_notice = None
        st.session_state.project_summary_cache = {}
        st.session_state.workflow_queue = []
        st.session_state.workflow_name = None
        st.session_state.workflow_notice = None
        st.session_state.task_board_cache = None
        st.session_state.daily_summary_cache = None
        st.rerun()


# =========================================================
# Work Mode + Task Template
# =========================================================
mode_col, task_col = st.columns(2)

with mode_col:
    previous_mode = st.session_state.work_mode
    st.selectbox("工作模式", all_mode_options(), key="work_mode", label_visibility="collapsed")
    if st.session_state.work_mode != previous_mode:
        ensure_valid_task_template()

with task_col:
    ensure_valid_task_template()
    st.selectbox("任務模板", get_task_options(st.session_state.work_mode), key="task_template", label_visibility="collapsed")

st.markdown(f'<div class="lobster-note">目前模式：{st.session_state.work_mode} · {mode_help_text(st.session_state.work_mode)}</div>', unsafe_allow_html=True)
st.markdown(f'<div class="lobster-note">目前任務：{st.session_state.task_template} · {task_help_text(st.session_state.work_mode, st.session_state.task_template)}</div>', unsafe_allow_html=True)

if st.session_state.auto_memory_notice:
    st.markdown(f'<div class="lobster-note">🧠 已自動記住：{st.session_state.auto_memory_notice}</div>', unsafe_allow_html=True)

if st.session_state.workflow_notice:
    st.markdown(f'<div class="lobster-note">⚙️ {st.session_state.workflow_notice}</div>', unsafe_allow_html=True)


# =========================================================
# Tool Switcher
# =========================================================
tool_panel = st.radio(
    "工具列",
    ["Desk", "Daily", "Board", "Outputs", "Chats", "Upload", "Memory"],
    horizontal=True,
    key="tool_panel",
    label_visibility="collapsed",
)

# =========================================================
# Desk
# =========================================================
if tool_panel == "Desk":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">🗂️ 專案工作台</div>', unsafe_allow_html=True)

    d1, d2, d3 = st.columns(3)
    with d1:
        st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
        st.markdown('<div class="lobster-mini-title">目前模式</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="lobster-muted">{st.session_state.work_mode}<br>{st.session_state.task_template}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    with d2:
        st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
        st.markdown('<div class="lobster-mini-title">目前檔案</div>', unsafe_allow_html=True)
        current_file = st.session_state.uploaded_file_cache["name"] if st.session_state.uploaded_file_cache else "尚未上傳"
        st.markdown(f'<div class="lobster-muted">{current_file}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    with d3:
        st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
        st.markdown('<div class="lobster-mini-title">目前對話</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="lobster-muted">訊息數：{len(st.session_state.messages)}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="lobster-mini-title">今日工作快覽</div>', unsafe_allow_html=True)
    daily = build_daily_summary(force_refresh=False)
    dd1, dd2, dd3, dd4 = st.columns(4)
    with dd1:
        st.markdown(f"**今日新任務**  \n{daily['new_tasks_count']} 筆")
    with dd2:
        st.markdown(f"**今日完成/交付**  \n{daily['done_tasks_count']} 筆")
    with dd3:
        st.markdown(f"**今日成果**  \n{daily['outputs_count']} 筆")
    with dd4:
        st.markdown(f"**仍卡住**  \n{daily['blocked_count']} 筆")

    st.markdown('<div class="lobster-mini-title">專案摘要</div>', unsafe_allow_html=True)
    ps1, ps2 = st.columns([5, 1])
    with ps2:
        if st.button("更新摘要", use_container_width=True):
            refreshed = get_or_refresh_project_summary(force_refresh=True)
            st.session_state.daily_summary_cache = None
            if refreshed:
                st.success("摘要已更新")
            else:
                st.warning("目前無法產生摘要")
            st.rerun()

    project_summary = get_or_refresh_project_summary(force_refresh=False) if st.session_state.messages else None
    if project_summary:
        st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
        st.markdown(f"**{project_summary.get('project_name', '未命名專案')}**")
        st.markdown(f"**目前重點**：{project_summary.get('current_focus', '')}")
        st.markdown(f"**整體狀態**：{project_summary.get('status_note', '')}")
        completed = project_summary.get("completed", []) or []
        missing = project_summary.get("missing", []) or []
        next_steps = project_summary.get("next_steps", []) or []
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("**已完成**")
            for x in completed[:5]:
                st.markdown(f"- {x}")
        with c2:
            st.markdown("**還缺什麼**")
            for x in missing[:5]:
                st.markdown(f"- {x}")
        with c3:
            st.markdown("**建議下一步**")
            for x in next_steps[:5]:
                st.markdown(f"- {x}")
        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.caption("目前摘要資料不足，先做幾輪工作後會更有意義。")

    st.markdown('<div class="lobster-mini-title">任務看板摘要</div>', unsafe_allow_html=True)
    board_tasks = load_task_board(force_refresh=False)
    status_counts = {status: 0 for status in TASK_STATUS_OPTIONS}
    total_linked_outputs = 0
    for t in board_tasks:
        status_counts[t["status"]] += 1
        total_linked_outputs += len(t["linked_outputs"])

    b1, b2, b3, b4 = st.columns(4)
    with b1:
        st.markdown(f"**待處理**  \n{status_counts['待處理']} 筆")
    with b2:
        st.markdown(f"**進行中**  \n{status_counts['進行中']} 筆")
    with b3:
        st.markdown(f"**已完成**  \n{status_counts['已完成']} 筆")
    with b4:
        st.markdown(f"**已交付**  \n{status_counts['已交付']} 筆")
    st.caption(f"目前任務已掛成果：{total_linked_outputs} 筆")

    st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">內建流程模板</div>', unsafe_allow_html=True)
    builtin_names = list(BUILTIN_WORKFLOW_TEMPLATES.keys())
    if builtin_names:
        cols = st.columns(min(3, len(builtin_names)))
        for i, name in enumerate(builtin_names):
            with cols[i % len(cols)]:
                if st.button(name, key=f"builtin_{name}", use_container_width=True):
                    start_workflow(name)

    st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">自訂流程模板</div>', unsafe_allow_html=True)
    custom_workflows = load_custom_workflows_from_memory(force_refresh=False)
    if custom_workflows:
        custom_names = list(custom_workflows.keys())
        cols = st.columns(min(3, len(custom_names)))
        for i, name in enumerate(custom_names):
            with cols[i % len(cols)]:
                if st.button(name, key=f"custom_run_{name}", use_container_width=True):
                    start_workflow(name)

        with st.expander("管理自訂流程", expanded=False):
            selected_delete = st.selectbox("選擇要刪除的流程", custom_names, key=f"delete_workflow_select_{st.session_state.workflow_refresh_key}")
            if st.button("刪除自訂流程", use_container_width=True):
                delete_custom_workflow(selected_delete)
                st.session_state.workflow_refresh_key += 1
                st.success("已刪除自訂流程")
                st.rerun()
    else:
        st.caption("目前沒有自訂流程。")

    with st.expander("新增自訂流程", expanded=False):
        wf_name = st.text_input("流程名稱", key=f"wf_name_{st.session_state.workflow_refresh_key}")
        steps_data = []
        step_count = st.selectbox("步驟數量", [1, 2, 3, 4, 5], index=2, key=f"wf_step_count_{st.session_state.workflow_refresh_key}")

        for i in range(step_count):
            st.markdown(f"**步驟 {i+1}**")
            col1, col2 = st.columns(2)
            with col1:
                step_mode = st.selectbox(f"模式 {i+1}", all_mode_options(), key=f"wf_mode_{i}_{st.session_state.workflow_refresh_key}")
            with col2:
                step_template = st.selectbox(f"任務模板 {i+1}", list(TASK_TEMPLATES[step_mode].keys()), key=f"wf_template_{i}_{st.session_state.workflow_refresh_key}")
            step_prompt = st.text_area(f"Prompt {i+1}", height=80, key=f"wf_prompt_{i}_{st.session_state.workflow_refresh_key}", placeholder="例如：請整理目前上傳檔案並列出 5 個重點")
            steps_data.append({"mode": step_mode, "template": step_template, "prompt": step_prompt.strip()})

        if st.button("保存自訂流程", use_container_width=True):
            clean_name = (wf_name or "").strip()
            clean_steps = [s for s in steps_data if s["prompt"]]
            if not clean_name:
                st.warning("請先輸入流程名稱。")
            elif not clean_steps:
                st.warning("至少要有一個有效步驟。")
            else:
                save_custom_workflow(clean_name, clean_steps)
                st.session_state.workflow_refresh_key += 1
                st.success("已保存自訂流程")
                st.rerun()

    st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">輸出格式化</div>', unsafe_allow_html=True)
    f1, f2, f3, f4 = st.columns(4)
    current_text = "\n\n".join([m["content"] for m in st.session_state.messages[-6:] if m["role"] == "assistant" and m.get("status") == "completed"]).strip()
    with f1:
        if st.button("轉主管版", use_container_width=True):
            if current_text:
                run_assistant_turn(format_output_prompt(current_text, "主管版"), "Report", "正式報告")
            else:
                st.warning("目前沒有可格式化的成果。")
    with f2:
        if st.button("轉客戶版", use_container_width=True):
            if current_text:
                run_assistant_turn(format_output_prompt(current_text, "客戶版"), "Report", "正式報告")
            else:
                st.warning("目前沒有可格式化的成果。")
    with f3:
        if st.button("轉簡報版", use_container_width=True):
            if current_text:
                run_assistant_turn(format_output_prompt(current_text, "簡報版"), "Report", "提案草稿")
            else:
                st.warning("目前沒有可格式化的成果。")
    with f4:
        if st.button("轉條列版", use_container_width=True):
            if current_text:
                run_assistant_turn(format_output_prompt(current_text, "條列版"), "Summary", "條列整理")
            else:
                st.warning("目前沒有可格式化的成果。")

    st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">成果下載中心</div>', unsafe_allow_html=True)
    latest_outputs = get_recent_output_cards(st.session_state.messages, max_cards=1)
    if latest_outputs:
        latest = latest_outputs[0]
        dl_payloads = get_download_payload(latest["assistant_full"], f"{latest['mode']}_{latest['template']}")
        dc1, dc2, dc3, dc4 = st.columns(4)
        with dc1:
            st.download_button(dl_payloads["txt"]["label"], data=dl_payloads["txt"]["data"], file_name=dl_payloads["txt"]["file_name"], mime=dl_payloads["txt"]["mime"], use_container_width=True, key="desk_download_txt")
        with dc2:
            st.download_button(dl_payloads["md"]["label"], data=dl_payloads["md"]["data"], file_name=dl_payloads["md"]["file_name"], mime=dl_payloads["md"]["mime"], use_container_width=True, key="desk_download_md")
        with dc3:
            st.download_button(dl_payloads["docx"]["label"], data=dl_payloads["docx"]["data"], file_name=dl_payloads["docx"]["file_name"], mime=dl_payloads["docx"]["mime"], use_container_width=True, key="desk_download_docx")
        with dc4:
            st.download_button(dl_payloads["pdf"]["label"], data=dl_payloads["pdf"]["data"], file_name=dl_payloads["pdf"]["file_name"], mime=dl_payloads["pdf"]["mime"], use_container_width=True, key="desk_download_pdf")
    else:
        st.caption("目前還沒有可下載的成果。")

    st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">快速工作</div>', unsafe_allow_html=True)
    w1, w2, w3 = st.columns(3)
    with w1:
        if st.button("摘要目前檔案", use_container_width=True):
            if st.session_state.uploaded_file_cache:
                st.session_state.work_mode = "Summary"
                st.session_state.task_template = "重點摘要"
                run_assistant_turn("請摘要目前上傳的檔案，先一句總結，再列 3 到 5 個重點，最後補充可執行建議。", "Summary", "重點摘要")
            else:
                st.warning("目前沒有檔案可摘要。")
    with w2:
        if st.button("分析目前檔案", use_container_width=True):
            if st.session_state.uploaded_file_cache:
                st.session_state.work_mode = "File Analyze"
                st.session_state.task_template = "結構分析"
                run_assistant_turn("請分析目前上傳的檔案，整理內容概覽、關鍵欄位或結構，並指出重點。", "File Analyze", "結構分析")
            else:
                st.warning("目前沒有檔案可分析。")
    with w3:
        if st.button("檢查異常", use_container_width=True):
            if st.session_state.uploaded_file_cache:
                st.session_state.work_mode = "File Analyze"
                st.session_state.task_template = "異常檢查"
                run_assistant_turn("請檢查目前上傳檔案中的異常、缺值、資料品質問題與風險。", "File Analyze", "異常檢查")
            else:
                st.warning("目前沒有檔案可檢查。")

    w4, w5, w6 = st.columns(3)
    with w4:
        if st.button("轉成正式報告", use_container_width=True):
            st.session_state.work_mode = "Report"
            st.session_state.task_template = "正式報告"
            if st.session_state.uploaded_file_cache:
                run_assistant_turn("請把目前上傳的檔案內容整理成正式報告。", "Report", "正式報告")
            elif st.session_state.messages:
                run_assistant_turn("請把目前這段對話的重點整理成正式報告。", "Report", "正式報告")
            else:
                st.warning("目前沒有可整理的內容。")
    with w5:
        if st.button("提煉成記憶", use_container_width=True):
            st.session_state.work_mode = "Memory Assist"
            st.session_state.task_template = "提煉長期記憶"
            if st.session_state.uploaded_file_cache:
                run_assistant_turn("請從目前上傳的檔案或上下文中，提煉值得長期保存的記憶。", "Memory Assist", "提煉長期記憶")
            elif st.session_state.messages:
                run_assistant_turn("請從目前這段對話中，提煉值得長期保存的記憶。", "Memory Assist", "提煉長期記憶")
            else:
                st.warning("目前沒有可提煉的內容。")
    with w6:
        if st.button("產出下一步建議", use_container_width=True):
            st.session_state.work_mode = "Chat"
            st.session_state.task_template = "行動建議"
            if st.session_state.uploaded_file_cache:
                run_assistant_turn("請根據目前上傳的檔案，產出清楚的下一步行動建議。", "Chat", "行動建議")
            elif st.session_state.messages:
                run_assistant_turn("請根據目前這段對話，產出清楚的下一步行動建議。", "Chat", "行動建議")
            else:
                st.warning("目前沒有足夠內容可給出下一步建議。")

    recent_outputs = get_recent_output_cards(st.session_state.messages, max_cards=6)
    if recent_outputs:
        st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">最近工作成果</div>', unsafe_allow_html=True)
        for item in recent_outputs[:6]:
            st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
            st.markdown(f'<span class="lobster-pill">{item["mode"]}</span><span class="lobster-pill">{item["template"]}</span>', unsafe_allow_html=True)
            st.markdown(f'<div class="lobster-muted"><strong>任務：</strong>{clip_preview(item["user_prompt"], 80)}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="lobster-muted" style="margin-top:6px;">{item["assistant_preview"]}</div>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

    suggestions = get_suggested_next_actions(st.session_state.messages, st.session_state.uploaded_file_cache is not None)
    if suggestions:
        st.markdown('<div class="lobster-mini-title" style="margin-top:0.25rem;">建議下一步</div>', unsafe_allow_html=True)
        for s in suggestions:
            st.markdown(f"- {s}")

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# Daily
# =========================================================
elif tool_panel == "Daily":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">📅 每日工作摘要</div>', unsafe_allow_html=True)

    top1, top2, top3, top4, top5 = st.columns(5)
    with top5:
        if st.button("刷新今日摘要", use_container_width=True):
            st.session_state.daily_summary_cache = None
            st.rerun()

    daily = build_daily_summary(force_refresh=False)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"**今日新任務**  \n{daily['new_tasks_count']} 筆")
    with c2:
        st.markdown(f"**今日完成/交付**  \n{daily['done_tasks_count']} 筆")
    with c3:
        st.markdown(f"**今日成果**  \n{daily['outputs_count']} 筆")
    with c4:
        st.markdown(f"**仍卡住**  \n{daily['blocked_count']} 筆")

    d1, d2 = st.columns(2)
    with d1:
        st.markdown("### 今天新增的任務")
        if daily["new_tasks"]:
            for x in daily["new_tasks"]:
                st.markdown(f"- {x}")
        else:
            st.caption("今天沒有新增任務")

        st.markdown("### 今天的成果")
        if daily["outputs"]:
            for x in daily["outputs"]:
                st.markdown(f"- {x}")
        else:
            st.caption("今天還沒有新成果")

    with d2:
        st.markdown("### 今天完成 / 交付")
        if daily["done_tasks"]:
            for x in daily["done_tasks"]:
                st.markdown(f"- {x}")
        else:
            st.caption("今天還沒有完成或交付的任務")

        st.markdown("### 目前仍卡住")
        if daily["blocked"]:
            for x in daily["blocked"]:
                st.markdown(f"- {x}")
        else:
            st.caption("目前沒有待處理或進行中的項目")

    st.markdown("### 建議下一步")
    if daily["next_steps"]:
        for x in daily["next_steps"]:
            st.markdown(f"- {x}")
    else:
        st.caption("目前沒有特別的下一步建議")

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# Board
# =========================================================
elif tool_panel == "Board":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">📋 任務看板</div>', unsafe_allow_html=True)

    with st.expander("新增任務", expanded=False):
        new_task_title = st.text_input("任務標題", key=f"new_task_title_{st.session_state.task_board_refresh_key}")
        new_task_desc = st.text_area("任務說明", height=100, key=f"new_task_desc_{st.session_state.task_board_refresh_key}")
        new_task_status = st.selectbox("初始狀態", TASK_STATUS_OPTIONS, key=f"new_task_status_{st.session_state.task_board_refresh_key}")
        if st.button("新增到看板", use_container_width=True):
            if new_task_title.strip():
                create_task_item(new_task_title, new_task_desc, new_task_status)
                st.session_state.task_board_refresh_key += 1
                st.session_state.daily_summary_cache = None
                st.success("已新增任務")
                st.rerun()
            else:
                st.warning("請輸入任務標題。")

    tasks = load_task_board(force_refresh=False)
    grouped = {status: [] for status in TASK_STATUS_OPTIONS}
    for task in tasks:
        grouped[task["status"]].append(task)

    cols = st.columns(4)
    for idx, status in enumerate(TASK_STATUS_OPTIONS):
        with cols[idx]:
            st.markdown('<div class="lobster-task-col">', unsafe_allow_html=True)
            st.markdown(f"**{status}**")
            st.markdown(f'<div class="lobster-task-count">{len(grouped[status])} 筆</div>', unsafe_allow_html=True)

            if not grouped[status]:
                st.caption("沒有任務")
            else:
                for task in grouped[status]:
                    st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
                    st.markdown(f"**{task['title']}**")
                    if task["description"]:
                        st.markdown(task["description"])

                    if task["linked_outputs"]:
                        st.markdown(f"**關聯成果：{len(task['linked_outputs'])} 筆**")
                        for i, linked in enumerate(task["linked_outputs"], start=1):
                            st.markdown('<div class="lobster-linked-box">', unsafe_allow_html=True)
                            st.markdown(f"**{i}. {linked['label']}**")
                            if linked.get("preview"):
                                st.markdown(linked["preview"])
                            with st.expander("看成果全文", expanded=False):
                                st.markdown(linked.get("content", ""))
                            st.markdown("</div>", unsafe_allow_html=True)

                    with st.expander("編輯任務", expanded=False):
                        edit_title = st.text_input("標題", value=task["title"], key=f"task_title_{task['memory_id']}")
                        edit_desc = st.text_area("說明", value=task["description"], height=100, key=f"task_desc_{task['memory_id']}")
                        edit_status = st.selectbox("狀態", TASK_STATUS_OPTIONS, index=TASK_STATUS_OPTIONS.index(task["status"]), key=f"task_status_{task['memory_id']}")

                        t1, t2 = st.columns(2)
                        with t1:
                            if st.button("保存", key=f"task_save_{task['memory_id']}", use_container_width=True):
                                update_task_item(
                                    task["memory_id"],
                                    edit_title,
                                    edit_desc,
                                    edit_status,
                                    linked_outputs=task["linked_outputs"],
                                )
                                st.session_state.daily_summary_cache = None
                                st.success("已更新任務")
                                st.rerun()
                        with t2:
                            if st.button("刪除", key=f"task_delete_{task['memory_id']}", use_container_width=True):
                                delete_task_item(task["memory_id"])
                                st.session_state.daily_summary_cache = None
                                st.success("已刪除任務")
                                st.rerun()

                    st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# Outputs
# =========================================================
elif tool_panel == "Outputs":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">📦 輸出物中心</div>', unsafe_allow_html=True)

    outputs = get_output_records(st.session_state.messages)
    if not outputs:
        st.caption("目前還沒有可管理的輸出物。先做一次摘要、分析或報告。")
    else:
        filter_col1, filter_col2 = st.columns(2)
        with filter_col1:
            mode_filter = st.selectbox("篩選模式", ["全部", "Chat", "Summary", "File Analyze", "Report", "Memory Assist"], key="outputs_mode_filter")
        with filter_col2:
            template_options = ["全部"] + sorted(list({o["template"] for o in outputs}))
            template_filter = st.selectbox("篩選任務", template_options, key="outputs_template_filter")

        filtered = []
        for o in outputs:
            if mode_filter != "全部" and o["mode"] != mode_filter:
                continue
            if template_filter != "全部" and o["template"] != template_filter:
                continue
            filtered.append(o)

        st.caption(f"目前顯示 {len(filtered)} 筆輸出物")
        existing_tasks = load_task_board(force_refresh=False)

        for o in filtered:
            st.markdown('<div class="lobster-soft-card">', unsafe_allow_html=True)
            st.markdown(f'<span class="lobster-pill">{o["mode"]}</span><span class="lobster-pill">{o["template"]}</span>', unsafe_allow_html=True)
            st.markdown(f"**原始任務**：{o['user_prompt']}")
            st.markdown(o["assistant_preview"])

            a1, a2, a3, a4 = st.columns(4)
            with a1:
                if st.button("再整理成報告", key=f"to_report_{o['id']}", use_container_width=True):
                    st.session_state.work_mode = "Report"
                    st.session_state.task_template = "正式報告"
                    run_assistant_turn(f"請把以下成果整理成正式報告：\n\n{o['assistant_full']}", "Report", "正式報告")
            with a2:
                if st.button("提煉成記憶", key=f"to_memory_{o['id']}", use_container_width=True):
                    create_memory_item(
                        title=f"{o['mode']}｜{shorten_text(o['user_prompt'], 40)}",
                        content=o["assistant_full"],
                        tags=[o["mode"], o["template"]],
                        source="output_center",
                    )
                    st.success("已存成記憶")
            with a3:
                if st.button("產出下一步", key=f"to_next_{o['id']}", use_container_width=True):
                    st.session_state.work_mode = "Chat"
                    st.session_state.task_template = "行動建議"
                    run_assistant_turn(f"請根據以下成果，產出清楚的下一步行動建議：\n\n{o['assistant_full']}", "Chat", "行動建議")
            with a4:
                if st.button("新任務＋掛成果", key=f"to_task_{o['id']}", use_container_width=True):
                    create_task_item(
                        title=shorten_text(o["user_prompt"], 60),
                        description=o["assistant_preview"],
                        status="待處理",
                        linked_outputs=[build_output_link_record(o)],
                    )
                    st.session_state.daily_summary_cache = None
                    st.success("已建立任務並掛上成果")

            if existing_tasks:
                st.markdown("**掛到既有任務**")
                link_col1, link_col2 = st.columns([3, 1])
                with link_col1:
                    task_options = {f"{t['title']}（{t['status']}）": t["memory_id"] for t in existing_tasks}
                    selected_task_label = st.selectbox(
                        "選擇任務",
                        list(task_options.keys()),
                        key=f"link_task_select_{o['id']}",
                        label_visibility="collapsed",
                    )
                with link_col2:
                    if st.button("掛上", key=f"link_task_btn_{o['id']}", use_container_width=True):
                        add_output_to_existing_task(task_options[selected_task_label], o)
                        st.session_state.daily_summary_cache = None
                        st.success("已掛到既有任務")
                        st.rerun()

            f1, f2, f3, f4 = st.columns(4)
            with f1:
                if st.button("主管版", key=f"to_exec_{o['id']}", use_container_width=True):
                    run_assistant_turn(format_output_prompt(o["assistant_full"], "主管版"), "Report", "正式報告")
            with f2:
                if st.button("客戶版", key=f"to_client_{o['id']}", use_container_width=True):
                    run_assistant_turn(format_output_prompt(o["assistant_full"], "客戶版"), "Report", "正式報告")
            with f3:
                if st.button("簡報版", key=f"to_slide_{o['id']}", use_container_width=True):
                    run_assistant_turn(format_output_prompt(o["assistant_full"], "簡報版"), "Report", "提案草稿")
            with f4:
                if st.button("條列版", key=f"to_bullet_{o['id']}", use_container_width=True):
                    run_assistant_turn(format_output_prompt(o["assistant_full"], "條列版"), "Summary", "條列整理")

            st.markdown("**下載成果**")
            payloads = get_download_payload(o["assistant_full"], f"{o['mode']}_{o['template']}_{o['id']}")
            d1, d2, d3, d4 = st.columns(4)
            with d1:
                st.download_button(payloads["txt"]["label"], data=payloads["txt"]["data"], file_name=payloads["txt"]["file_name"], mime=payloads["txt"]["mime"], use_container_width=True, key=f"dl_txt_{o['id']}")
            with d2:
                st.download_button(payloads["md"]["label"], data=payloads["md"]["data"], file_name=payloads["md"]["file_name"], mime=payloads["md"]["mime"], use_container_width=True, key=f"dl_md_{o['id']}")
            with d3:
                st.download_button(payloads["docx"]["label"], data=payloads["docx"]["data"], file_name=payloads["docx"]["file_name"], mime=payloads["docx"]["mime"], use_container_width=True, key=f"dl_docx_{o['id']}")
            with d4:
                st.download_button(payloads["pdf"]["label"], data=payloads["pdf"]["data"], file_name=payloads["pdf"]["file_name"], mime=payloads["pdf"]["mime"], use_container_width=True, key=f"dl_pdf_{o['id']}")

            with st.expander("看完整內容", expanded=False):
                st.markdown(o["assistant_full"])
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# Chats
# =========================================================
elif tool_panel == "Chats":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">💬 Chats</div>', unsafe_allow_html=True)

    sessions = list_sessions()
    if not sessions:
        st.caption("目前沒有聊天紀錄")
    else:
        session_options = []
        session_map = {}
        for s in sessions[:50]:
            label = s["title"] or "New Chat"
            display = f"● {label}" if s["id"] == st.session_state.active_chat_id else label
            session_options.append(display)
            session_map[display] = s["id"]

        current_display = None
        for k, v in session_map.items():
            if v == st.session_state.active_chat_id:
                current_display = k
                break

        selected_display = st.selectbox(
            "選擇對話",
            options=session_options,
            index=session_options.index(current_display) if current_display in session_options else 0,
            label_visibility="collapsed",
        )

        selected_session_id = session_map[selected_display]
        if selected_session_id != st.session_state.active_chat_id:
            st.session_state.active_chat_id = selected_session_id
            load_messages(selected_session_id)
            st.session_state.last_memory_suggestion = None
            st.session_state.auto_memory_notice = None
            st.session_state.workflow_queue = []
            st.session_state.workflow_name = None
            st.session_state.workflow_notice = None
            st.session_state.daily_summary_cache = None
            st.rerun()

        st.selectbox("訊息顯示數量", ["最近 20 則", "最近 30 則", "最近 50 則", "全部顯示"], key="message_window", label_visibility="collapsed")

        with st.expander("管理目前對話", expanded=False):
            active_title = next((s["title"] for s in sessions if s["id"] == st.session_state.active_chat_id), "New Chat")
            new_title = st.text_input("重新命名", value=active_title, key="active_chat_rename_input", placeholder="輸入對話名稱")

            c1, c2 = st.columns(2)
            with c1:
                if st.button("保存名稱", use_container_width=True):
                    rename_session(st.session_state.active_chat_id, new_title)
                    st.rerun()
            with c2:
                if st.button("刪除對話", use_container_width=True):
                    delete_session(st.session_state.active_chat_id)
                    remaining = list_sessions()
                    if remaining:
                        st.session_state.active_chat_id = remaining[0]["id"]
                        load_messages(remaining[0]["id"])
                    else:
                        new_id = create_session("New Chat")
                        st.session_state.active_chat_id = new_id
                        st.session_state.messages = []
                    st.session_state.last_memory_suggestion = None
                    st.session_state.auto_memory_notice = None
                    st.session_state.workflow_queue = []
                    st.session_state.workflow_name = None
                    st.session_state.workflow_notice = None
                    st.session_state.daily_summary_cache = None
                    st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# Upload
# =========================================================
elif tool_panel == "Upload":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">📎 Upload</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Upload context",
        type=["txt", "pdf", "csv", "xlsx", "xls", "png", "jpg", "jpeg"],
        key="context_file",
        label_visibility="collapsed",
    )

    if uploaded_file is not None:
        st.session_state.uploaded_file_cache = {
            "name": uploaded_file.name,
            "type": uploaded_file.type,
            "data": uploaded_file.getvalue(),
        }
        st.success(f"已載入：{uploaded_file.name}")

    if st.session_state.uploaded_file_cache:
        st.caption(f"目前檔案：{st.session_state.uploaded_file_cache['name']}")
        if st.button("清除檔案", use_container_width=True):
            st.session_state.uploaded_file_cache = None
            st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# Memory
# =========================================================
elif tool_panel == "Memory":
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">🧠 Memory</div>', unsafe_allow_html=True)

    with st.expander("新增記憶", expanded=False):
        mem_title = st.text_input("記憶標題", key=f"mem_title_{st.session_state.memory_refresh_key}")
        mem_tags = st.text_input("標籤（逗號分隔）", key=f"mem_tags_{st.session_state.memory_refresh_key}")
        mem_content = st.text_area("記憶內容", height=96, key=f"mem_content_{st.session_state.memory_refresh_key}")

        if st.button("保存記憶", use_container_width=True):
            if mem_content.strip():
                tags = [t.strip() for t in mem_tags.split(",")] if mem_tags.strip() else []
                create_memory_item(title=mem_title or mem_content[:30], content=mem_content, tags=tags, source="manual")
                st.session_state.memory_refresh_key += 1
                st.rerun()

    memory_items = list_memory_items(limit=5)
    if not memory_items:
        st.caption("目前沒有長期記憶")
    else:
        for item in memory_items:
            st.markdown(
                f"""
                <div class="lobster-soft-card">
                    <div class="lobster-mini-title">{item.get('title', '')}</div>
                    <div class="lobster-muted">{", ".join(item.get('tags', []) or [])}</div>
                    <div class="lobster-muted">{(item.get('content', '')[:120] + '...') if len(item.get('content', '')) > 120 else item.get('content', '')}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if st.button("刪除這筆記憶", key=f"del_mem_{item['id']}", use_container_width=True):
                delete_memory_item(item["id"])
                st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)


# =========================================================
# Memory Assist Save Panel
# =========================================================
if st.session_state.work_mode == "Memory Assist" and st.session_state.last_memory_suggestion:
    suggestion = st.session_state.last_memory_suggestion
    st.markdown('<div class="lobster-card">', unsafe_allow_html=True)
    st.markdown('<div class="lobster-section-title">📝 記憶建議</div>', unsafe_allow_html=True)

    st.caption(f"是否值得記住：{suggestion.get('worth', '')}")
    if suggestion.get("reason"):
        st.caption(f"理由：{suggestion.get('reason', '')}")

    save_title = st.text_input("建議標題", value=suggestion.get("title", ""), key="memory_assist_title")
    save_tags = st.text_input("建議標籤（逗號分隔）", value=", ".join(suggestion.get("tags", [])), key="memory_assist_tags")
    save_content = st.text_area("建議記憶內容", value=suggestion.get("content", ""), height=120, key="memory_assist_content")

    cc1, cc2 = st.columns(2)
    with cc1:
        if st.button("存成記憶", use_container_width=True):
            if save_content.strip():
                tags = [t.strip() for t in save_tags.split(",")] if save_tags.strip() else []
                create_memory_item(title=save_title or save_content[:30], content=save_content, tags=tags, source="memory_assist")
                st.success("已存成記憶")
                st.session_state.memory_refresh_key += 1
    with cc2:
        if st.button("清除建議", use_container_width=True):
            st.session_state.last_memory_suggestion = None
            st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)


# =========================================================
# Main Chat Area
# =========================================================
all_messages = st.session_state.messages

if not all_messages:
    st.markdown(
        """
        <div class="lobster-empty">
            <h2>今天想讓龍蝦做什麼？</h2>
            <div>先選工作模式與任務模板，或直接使用工作台、每日摘要、任務看板、流程模板、成果下載中心。</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
else:
    total_count = len(all_messages)
    visible_count = get_window_size(st.session_state.message_window, total_count)

    if visible_count < total_count:
        hidden_count = total_count - visible_count
        st.markdown(f'<div class="lobster-note">已隱藏較早的 {hidden_count} 則訊息。</div>', unsafe_allow_html=True)

    messages_to_show = all_messages[-visible_count:]
    for msg in messages_to_show:
        display_content = msg["content"]
        if msg["role"] == "assistant":
            if msg.get("status") == "pending":
                display_content = "🦞 龍蝦正在思考中..."
            elif msg.get("status") == "failed" and msg.get("error_message"):
                display_content = f"{msg['content']}\n\n> {msg['error_message']}"

        with st.chat_message("assistant" if msg["role"] == "assistant" else "user"):
            st.markdown(display_content)


# =========================================================
# Chat Input
# =========================================================
prompt_placeholder_map = {
    "Chat": "輸入問題、想法或要改寫的內容...",
    "Summary": "輸入要摘要的內容，或先上傳檔案...",
    "File Analyze": "輸入分析需求，或先上傳檔案...",
    "Report": "輸入要整理成報告的主題或內容...",
    "Memory Assist": "輸入想整理成記憶的內容...",
}

prompt = st.chat_input(prompt_placeholder_map.get(st.session_state.work_mode, "Message Lobster..."))

if prompt:
    run_assistant_turn(prompt, st.session_state.work_mode, st.session_state.task_template)
